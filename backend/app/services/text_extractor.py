"""Local text extraction for PDFs, images, and text documents."""

from io import BytesIO
import shutil

from app.models.schemas import ExtractionError, ExtractionResult, UploadedDocument
from app.services.extraction_quality import assess_extraction_quality
from app.utils.text_cleaning import normalize_extracted_text, normalize_ocr_artifacts


class TextExtractionUnavailableError(RuntimeError):
    """Raised when a local extraction dependency is unavailable."""


def extract_text(document: UploadedDocument) -> ExtractionResult:
    """Extract clean text from a validated uploaded document."""
    extension = document.extension.lower()
    if extension in {".txt", ".md"}:
        text = _extract_plain_text(document.content)
        return _result(document, text=text, engine="plain_text")
    if extension == ".docx":
        return _extract_docx(document)
    if extension == ".pdf":
        return _extract_pdf(document)
    if extension in {".png", ".jpg", ".jpeg"}:
        return _extract_image(document)

    return _result(
        document,
        text="",
        engine="unsupported",
        error=ExtractionError(
            code="unsupported_type",
            message=f"Extraction is not implemented for {extension}.",
        ),
    )


def _extract_plain_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return normalize_extracted_text(content.decode(encoding))
        except UnicodeDecodeError:
            continue
    return normalize_extracted_text(content.decode("utf-8", errors="replace"))


def _extract_pdf(document: UploadedDocument) -> ExtractionResult:
    warnings: list[str] = []
    try:
        import fitz
    except ImportError:
        fitz = None
        warnings.append("PyMuPDF is not installed; trying pdfplumber fallback.")

    if fitz is not None:
        try:
            with fitz.open(stream=document.content, filetype="pdf") as pdf:
                pages = [page.get_text("text") for page in pdf]
                sparse_page_indexes = [
                    index
                    for index, page_text in enumerate(pages)
                    if len(page_text.strip()) < 20
                ]
                if sparse_page_indexes:
                    (
                        pages,
                        ocr_warnings,
                        ocr_engine,
                        ocr_confidence,
                        ocr_page_count,
                        ocr_error,
                    ) = _ocr_sparse_pdf_pages(
                        document=document,
                        pdf=pdf,
                        fitz=fitz,
                        pages=pages,
                        sparse_page_indexes=sparse_page_indexes,
                    )
                    warnings.extend(ocr_warnings)
                    if ocr_page_count:
                        warnings.append(
                            f"Used local OCR for {ocr_page_count} of "
                            f"{pdf.page_count} PDF pages with little or no embedded text."
                        )
                        return _result(
                            document,
                            text=normalize_extracted_text("\n\n".join(pages)),
                            warnings=warnings,
                            page_count=pdf.page_count,
                            engine=f"pymupdf+{ocr_engine or 'ocr'}",
                            is_ocr=True,
                            ocr_engine=ocr_engine,
                            ocr_mean_confidence=ocr_confidence,
                            page_text_coverage=_page_text_coverage(pages),
                        )
                    if not any(page.strip() for page in pages) and ocr_error:
                        return _result(
                            document,
                            text="",
                            warnings=warnings,
                            page_count=pdf.page_count,
                            engine="pymupdf+ocr",
                            is_ocr=True,
                            error=ExtractionError(
                                code="pdf_ocr_failed",
                                message=(
                                    "The PDF contained no embedded text and local OCR "
                                    f"could not recover readable content: {ocr_error.message}"
                                ),
                            ),
                            page_text_coverage=0.0,
                        )
                return _result(
                    document,
                    text=normalize_extracted_text("\n\n".join(pages)),
                    warnings=warnings,
                    page_count=pdf.page_count,
                    engine="pymupdf",
                    page_text_coverage=_page_text_coverage(pages),
                )
        except Exception as exc:
            warnings.append(f"PyMuPDF could not extract this PDF: {exc}")

    try:
        import pdfplumber
    except ImportError:
        return _result(
            document,
            text="",
            warnings=warnings,
            engine="pdf_unavailable",
            error=ExtractionError(
                code="pdf_dependency_missing",
                message="PDF extraction requires PyMuPDF or pdfplumber to be installed.",
            ),
        )

    try:
        with pdfplumber.open(BytesIO(document.content)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
            return _result(
                document,
                text=normalize_extracted_text("\n\n".join(pages)),
                warnings=warnings,
                page_count=len(pdf.pages),
                engine="pdfplumber",
                page_text_coverage=_page_text_coverage(pages),
            )
    except Exception as exc:
        return _result(
            document,
            text="",
            warnings=warnings,
            engine="pdfplumber",
            error=ExtractionError(
                code="pdf_extraction_failed",
                message=f"Could not extract text from the PDF: {exc}",
            ),
        )


def _ocr_sparse_pdf_pages(
    *,
    document: UploadedDocument,
    pdf: object,
    fitz: object,
    pages: list[str],
    sparse_page_indexes: list[int],
) -> tuple[
    list[str],
    list[str],
    str | None,
    float | None,
    int,
    ExtractionError | None,
]:
    warnings: list[str] = []
    ocr_engines: set[str] = set()
    confidences: list[float] = []
    ocr_page_count = 0
    first_error: ExtractionError | None = None

    for page_index in sparse_page_indexes:
        page = pdf[page_index]  # type: ignore[index]
        pixmap = page.get_pixmap(  # type: ignore[attr-defined]
            matrix=fitz.Matrix(2, 2),  # type: ignore[attr-defined]
            alpha=False,
        )
        image_bytes = pixmap.tobytes("png")
        page_document = UploadedDocument(
            filename=f"{document.filename}-page-{page_index + 1}.png",
            content=image_bytes,
            extension=".png",
            content_type="image/png",
            size_bytes=len(image_bytes),
        )
        page_result = _extract_image(page_document)
        for warning in page_result.warnings:
            if warning not in warnings:
                warnings.append(warning)
        if not page_result.text.strip():
            if first_error is None and page_result.error is not None:
                first_error = page_result.error
            continue

        pages[page_index] = page_result.text
        ocr_page_count += 1
        if page_result.diagnostics.ocr_engine:
            ocr_engines.add(page_result.diagnostics.ocr_engine)
        if page_result.diagnostics.ocr_mean_confidence is not None:
            confidences.append(page_result.diagnostics.ocr_mean_confidence)

    if len(ocr_engines) == 1:
        ocr_engine = next(iter(ocr_engines))
    elif ocr_engines:
        ocr_engine = "mixed"
    else:
        ocr_engine = None
    return (
        pages,
        warnings,
        ocr_engine,
        _mean(confidences),
        ocr_page_count,
        first_error,
    )


def _extract_docx(document: UploadedDocument) -> ExtractionResult:
    try:
        from docx import Document
    except ImportError:
        return _result(
            document,
            text="",
            engine="python_docx",
            error=ExtractionError(
                code="docx_dependency_missing",
                message="DOCX extraction requires python-docx to be installed.",
            ),
        )

    try:
        doc = Document(BytesIO(document.content))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        return _result(
            document,
            text=normalize_extracted_text("\n".join(paragraphs)),
            engine="python_docx",
        )
    except Exception as exc:
        return _result(
            document,
            text="",
            engine="python_docx",
            error=ExtractionError(
                code="docx_extraction_failed",
                message=f"Could not extract text from the DOCX file: {exc}",
            ),
        )


def _extract_image(document: UploadedDocument) -> ExtractionResult:
    if not _tesseract_is_available():
        return _extract_image_with_rapidocr(document)

    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        message = "Image OCR requires Pillow and pytesseract to be installed."
        return _result(
            document,
            text="",
            warnings=[message],
            engine="tesseract",
            is_ocr=True,
            ocr_engine="tesseract",
            error=ExtractionError(code="ocr_dependency_missing", message=message),
        )

    try:
        image = Image.open(BytesIO(document.content))
        text, mean_confidence = _tesseract_text_and_confidence(image, pytesseract)
        return _result(
            document,
            text=normalize_ocr_artifacts(text),
            engine="tesseract",
            is_ocr=True,
            ocr_engine="tesseract",
            ocr_mean_confidence=mean_confidence,
        )
    except Exception as exc:
        return _result(
            document,
            text="",
            engine="tesseract",
            is_ocr=True,
            ocr_engine="tesseract",
            error=ExtractionError(
                code="ocr_extraction_failed",
                message=f"Could not extract text from the image: {exc}",
            ),
        )


def _extract_image_with_rapidocr(document: UploadedDocument) -> ExtractionResult:
    """Extract image text with a local Python OCR fallback when available."""
    try:
        import numpy as np
        from PIL import Image
        from rapidocr_onnxruntime import RapidOCR
    except ImportError:
        message = (
            "No local OCR engine is available. Install Tesseract, or install the "
            "rapidocr-onnxruntime Python package for local image OCR without a "
            "separate system binary."
        )
        return _result(
            document,
            text="",
            warnings=[message],
            engine="rapidocr",
            is_ocr=True,
            ocr_engine="rapidocr",
            error=ExtractionError(code="ocr_unavailable", message=message),
        )

    try:
        image = Image.open(BytesIO(document.content)).convert("RGB")
        engine = RapidOCR()
        results, _ = engine(np.array(image))
        text_lines = [str(item[1]) for item in results or [] if len(item) >= 2]
        confidences = [
            float(item[2])
            for item in results or []
            if len(item) >= 3 and _is_number(item[2])
        ]
        text = normalize_ocr_artifacts("\n".join(text_lines))
        warnings = ["Tesseract was not found; used RapidOCR local fallback."]
        if not text.strip():
            return _result(
                document,
                text="",
                warnings=warnings,
                engine="rapidocr",
                is_ocr=True,
                ocr_engine="rapidocr",
                error=ExtractionError(
                    code="ocr_no_text",
                    message="OCR completed, but no readable text was found in this image.",
                ),
            )
        return _result(
            document,
            text=text,
            warnings=warnings,
            engine="rapidocr",
            is_ocr=True,
            ocr_engine="rapidocr",
            ocr_mean_confidence=_mean(confidences),
        )
    except Exception as exc:
        return _result(
            document,
            text="",
            engine="rapidocr",
            is_ocr=True,
            ocr_engine="rapidocr",
            error=ExtractionError(
                code="ocr_extraction_failed",
                message=f"Could not extract text from the image with local OCR: {exc}",
            ),
        )


def _tesseract_is_available() -> bool:
    """Return whether the local Tesseract binary appears available."""
    return shutil.which("tesseract") is not None


def _result(
    document: UploadedDocument,
    *,
    text: str,
    warnings: list[str] | None = None,
    error: ExtractionError | None = None,
    page_count: int | None = None,
    engine: str = "unknown",
    is_ocr: bool = False,
    ocr_engine: str | None = None,
    ocr_mean_confidence: float | None = None,
    page_text_coverage: float | None = None,
) -> ExtractionResult:
    diagnostics = assess_extraction_quality(
        text,
        engine=engine,
        is_ocr=is_ocr,
        ocr_engine=ocr_engine,
        ocr_mean_confidence=ocr_mean_confidence,
        page_text_coverage=page_text_coverage,
    )
    result_warnings = list(warnings or [])
    if diagnostics.confidence == "low" and text.strip():
        result_warnings.append(
            "Extraction quality is low; verify important details against the source."
        )
    return ExtractionResult(
        filename=document.filename,
        document_type=document.extension.removeprefix("."),
        text=text,
        warnings=result_warnings,
        error=error,
        page_count=page_count,
        metadata={
            "content_type": document.content_type,
            "size_bytes": document.size_bytes,
            "stored_original": False,
            "engine": engine,
        },
        diagnostics=diagnostics,
    )


def _page_text_coverage(pages: list[str]) -> float | None:
    if not pages:
        return None
    meaningful = sum(len(page.strip()) >= 20 for page in pages)
    return meaningful / len(pages)


def _tesseract_text_and_confidence(image: object, pytesseract: object) -> tuple[str, float | None]:
    output = pytesseract.image_to_data(  # type: ignore[attr-defined]
        image,
        output_type=pytesseract.Output.DICT,  # type: ignore[attr-defined]
    )
    lines: dict[tuple[int, int, int], list[str]] = {}
    confidences: list[float] = []
    texts = output.get("text", [])
    for index, token in enumerate(texts):
        cleaned = str(token).strip()
        if not cleaned:
            continue
        key = (
            int(output["block_num"][index]),
            int(output["par_num"][index]),
            int(output["line_num"][index]),
        )
        lines.setdefault(key, []).append(cleaned)
        confidence = output.get("conf", [])[index]
        if _is_number(confidence):
            normalized = float(confidence)
            if normalized >= 0:
                confidences.append(min(normalized / 100, 1.0))
    text = "\n".join(" ".join(tokens) for tokens in lines.values())
    return text, _mean(confidences)


def _is_number(value: object) -> bool:
    try:
        float(value)
    except (TypeError, ValueError):
        return False
    return True


def _mean(values: list[float]) -> float | None:
    if not values:
        return None
    return sum(values) / len(values)
